"""Create an editable two-column academic article draft for the project."""

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SRC_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SRC_DIR.parent
REPORT_DIR = PROJECT_DIR / "report"
ASSET_DIR = REPORT_DIR / "assets_generated"
OUTPUT_DIR = Path("/workspace/output/docx")
OUTPUT_PATH = OUTPUT_DIR / "Neural_Guided_Astar_Makale_Taslagi.docx"
sys.path.insert(0, str(SRC_DIR))

import generate_final_report as report_data  # noqa: E402


# narrative_proposal preset + named override: two-column academic paper
FONT = "Times New Roman"
BLUE = "1F4D78"
MUTED = "5D6670"
TABLE_FILL = "EEF2F5"
COLUMN_DXA = 4774


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_section(section, columns=1):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "432")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    heading_specs = {
        "Heading 1": (11, 8, 3),
        "Heading 2": (10, 6, 2),
        "Heading 3": (9.5, 4, 2),
    }
    for name, (size, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Makale Caption" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("Makale Caption", 1)
    else:
        caption = doc.styles["Makale Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(7.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(4)

    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(9)
    list_style.paragraph_format.left_indent = Inches(0.22)
    list_style.paragraph_format.first_line_indent = Inches(-0.12)
    list_style.paragraph_format.space_after = Pt(2)
    list_style.paragraph_format.line_spacing = 1.0


def configure_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_page_field(p)
    for run in p.runs:
        set_run_font(run, size=8, color=MUTED)


def set_cell_margins(cell, top=50, start=70, bottom=50, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "70")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc, rows, widths_dxa, caption):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = str(value)
            if r == 0:
                shade_cell(cell, TABLE_FILL)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=6.8, bold=(r == 0))
    set_table_geometry(table, widths_dxa)
    p = doc.add_paragraph(caption, style="Makale Caption")
    p.paragraph_format.keep_with_next = False
    return table


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        run.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_picture(doc, path, width_in, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in))
    doc.add_paragraph(caption, style="Makale Caption")


def add_title_block(doc):
    logo = REPORT_DIR / "biruni_logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(logo), width=Inches(0.65))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Dinamik Tehdit Ortamlarında İHA Rota Planlaması için\nNeural-Guided A* Yaklaşımı")
    set_run_font(run, size=16, bold=True, color="17365D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Kaan Arslan¹, Rukiye Narsu Oymak¹")
    set_run_font(run, size=10.5, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("¹Biruni Üniversitesi, Yapay Zeka Dersi Dönem Projesi\nDanışman: Mahyar Teymournezhad")
    set_run_font(run, size=8.5, italic=True, color=MUTED)


def add_abstract(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = p.add_run("Özet—")
    set_run_font(lead, size=9, bold=True)
    text = (
        "Bu çalışmada, statik engeller ve periyodik dinamik tehditler içeren iki boyutlu ortamlarda "
        "İHA rota planlaması için MLP tabanlı öğrenilmiş bir heuristic ile zaman-genişletmeli A* "
        "birleştirilmiştir. Dijkstra, Manhattan, Octile ve Neural yöntemleri dört senaryoda genişletilen "
        "düğüm ve rota maliyeti açısından karşılaştırılmıştır. Neural heuristic her düğümün varış zamanını "
        "kullanmakta; eğitim etiketleri yönlü hücre maliyetlerine uygun ters Dijkstra ile üretilmektedir. "
        "Sonuçlar Neural yaklaşımın Basit Sızma ve Şehir Operasyonu senaryolarında Octile'dan daha az düğüm "
        "genişlettiğini, Koridor ve Dinamik Tehdit senaryolarında ise Octile'ın daha verimli olduğunu göstermiştir."
    )
    run = p.add_run(text)
    set_run_font(run, size=9)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = p.add_run("Anahtar Kelimeler—")
    set_run_font(lead, size=8.5, bold=True)
    run = p.add_run("A* araması, Dijkstra, MLP, öğrenilmiş heuristic, İHA rota planlama, dinamik tehdit.")
    set_run_font(run, size=8.5, italic=True)


def build_article(results):
    doc = Document()
    configure_styles(doc)
    first = doc.sections[0]
    configure_section(first, columns=1)
    configure_footer(first)
    doc.core_properties.title = "Neural-Guided A* Yaklaşımı - Makale Taslağı"
    doc.core_properties.author = "Kaan Arslan ve Rukiye Narsu Oymak"
    add_title_block(doc)
    add_abstract(doc)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, columns=2)

    add_heading(doc, "1. Giriş")
    add_body(doc, "Otonom İHA sistemlerinde rota planlama, statik engellerin yanı sıra zamanla değişen radar ve SAM bölgelerinin de dikkate alınmasını gerektirir. Klasik A* algoritması uygun bir heuristic ile arama alanını daraltırken, öğrenilmiş heuristicler çevresel özellikleri kullanarak daha yönlendirici tahminler üretebilir. Bununla birlikte öğrenilmiş bir heuristic gerçek kalan maliyeti aşabildiği için optimalite garantisi taşımayabilir.")
    add_body(doc, "Bu çalışmanın amacı, zaman-genişletmeli A* içinde kullanılan MLP tabanlı Neural heuristicin klasik Dijkstra, Manhattan ve Octile yöntemleri karşısındaki performansını incelemektir. Değerlendirmede yalnızca genişletilen düğüm sayısı değil, rota maliyeti ve optimalite sapması da dikkate alınmıştır.")

    add_heading(doc, "2. Yöntem")
    add_heading(doc, "2.1 Operasyon Ortamı", 2)
    add_body(doc, "Ortam 50x50 hücreli iki boyutlu bir grid olarak modellenmiştir. İHA sekiz yönde hareket edebilmekte; düz hareketler 1, çapraz hareketler √2 maliyet taşımaktadır. Pasif tehdit hücrelerine 1,5 risk çarpanı uygulanmakta, aktif tehditler ve yasak bölgeler geçilemez kabul edilmektedir.")
    add_heading(doc, "2.2 Zaman-Genişletmeli A*", 2)
    add_body(doc, "Dinamik arama durumu (r,c,t) biçiminde tanımlanmıştır. Burada r satır, c sütun ve t zaman adımıdır. Her hareket veya bekleme eyleminde zaman ilerletilir ve hedef hücrenin güvenliği varış zamanında kontrol edilir. Tehdit periyotlarının EKOK'u kullanılarak aynı konumun farklı zaman fazları ayrı durumlar olarak ele alınır.")
    add_heading(doc, "2.3 Neural Heuristic", 2)
    add_body(doc, "MLP modeli sekiz giriş özelliği, 64-64-32 gizli katman ve tek lineer çıkıştan oluşmaktadır. Girdiler Manhattan, Euclidean ve Octile mesafeleri; hedef yönü; yerel engel yoğunluğu; hedef doğrultusundaki engel oranı ve hücre riskidir. Model Adam optimizer ve MSE kaybıyla eğitilmiştir.")
    add_picture(doc, ASSET_DIR / "training_curve.png", 3.12, "Şekil 1. MLP eğitim ve doğrulama kaybı.")
    add_body(doc, "Son model 157.693 eğitim ve 31.668 doğrulama örneğiyle 80 epoch eğitilmiştir. Eğitim ve doğrulama haritalarında farklı seed grupları kullanılmıştır.")

    add_heading(doc, "3. Deneysel Kurulum")
    add_body(doc, "Deneyler Basit Sızma, Şehir Operasyonu, Koridor Geçişi ve Dinamik Tehdit olmak üzere dört senaryoda yürütülmüştür. Dijkstra optimum referans, Octile güçlü admissible baseline, Manhattan ise sekiz yönlü modelde inadmissible karşılaştırma olarak kullanılmıştır.")

    node_rows = [["Senaryo", "Dij.", "Man.", "Oct.", "Neu."]]
    cost_rows = [["Senaryo", "Dij.", "Man.", "Oct.", "Neu."]]
    for scenario, title, _ in report_data.SCENARIOS:
        short_title = {"Basit Sizma": "Basit", "Sehir Operasyonu": "Şehir", "Koridor Gecisi": "Koridor", "Dinamik Tehdit": "Dinamik"}[title]
        rs = results[scenario]
        node_rows.append([
            short_title,
            f"{rs[report_data.METHODS[0]].nodes_expanded:,}".replace(",", "."),
            f"{rs[report_data.METHODS[1]].nodes_expanded:,}".replace(",", "."),
            f"{rs[report_data.METHODS[2]].nodes_expanded:,}".replace(",", "."),
            f"{rs[report_data.METHODS[3]].nodes_expanded:,}".replace(",", "."),
        ])
        cost_rows.append([
            short_title,
            f"{rs[report_data.METHODS[0]].cost:.2f}", f"{rs[report_data.METHODS[1]].cost:.2f}",
            f"{rs[report_data.METHODS[2]].cost:.2f}", f"{rs[report_data.METHODS[3]].cost:.2f}",
        ])
    add_table(doc, node_rows, [1400, 844, 844, 844, 842], "Tablo 1. Genişletilen düğüm sayıları.")
    add_table(doc, cost_rows, [1400, 844, 844, 844, 842], "Tablo 2. Rota maliyetleri.")

    add_heading(doc, "4. Bulgular")
    add_body(doc, "Neural yöntem Basit Sızma senaryosunda Octile'a göre %8,4, Şehir Operasyonu'nda %19,6 daha az düğüm genişletmiştir. Koridor senaryosunda %105,5 ve Dinamik Tehdit senaryosunda %15,0 daha fazla düğüm genişletmiştir. Neural rota maliyeti Dijkstra optimumuna göre sırasıyla %0,90, %2,03, %5,35 ve %0,91 sapmıştır.")

    figures = [
        ("01_basit", "Şekil 2. Basit Sızma senaryosu rota karşılaştırması."),
        ("02_sehir", "Şekil 3. Şehir Operasyonu rota karşılaştırması."),
        ("03_koridor", "Şekil 4. Koridor Geçişi rota karşılaştırması."),
        ("04_dinamik", "Şekil 5. Dinamik Tehdit rota karşılaştırması."),
    ]
    for prefix, caption in figures:
        add_picture(doc, ASSET_DIR / f"{prefix}_routes.png", 3.12, caption)

    add_heading(doc, "4.1 Heuristic Yüzeyleri", 2)
    add_body(doc, "Manhattan yüzeyi yalnızca geometrik uzaklığı yansıtırken Neural yüzey yerel engel ve risk özelliklerinden etkilenen daha düzensiz bir maliyet tahmini üretmektedir.")
    add_picture(doc, ASSET_DIR / "urban_manhattan_heatmap.png", 2.55, "Şekil 6. Manhattan heuristic heatmap.")
    add_picture(doc, ASSET_DIR / "urban_neural_heatmap.png", 2.55, "Şekil 7. Neural heuristic heatmap.")

    doc.add_page_break()
    add_heading(doc, "5. Tartışma")
    add_body(doc, "Deneyler Neural heuristicin performansının senaryo yapısına bağlı olduğunu göstermektedir. Basit ve şehir ortamlarında çevresel özellikler aramayı daha doğrudan yönlendirmiştir. Dar koridorda ise geometrik yapının belirgin olması nedeniyle Octile daha başarılıdır. Neural yöntem bazı senaryolarda daha az düğüm açsa da MLP çıkarım maliyeti ve optimalite sapması göz önünde bulundurulmalıdır.")
    add_body(doc, "Manhattan çok az düğüm genişletmesine rağmen sekiz yönlü hareket modelinde admissible değildir. Bu nedenle en az düğüm sayısı tek başına en iyi yöntem anlamına gelmemektedir. Mevcut deneyler açısından Octile genel olarak en dengeli ve güvenilir yöntemdir; Neural yaklaşım ise belirli ortam türlerinde arama verimliliği sağlayan tamamlayıcı bir yöntemdir.")

    add_heading(doc, "5.1 Sınırlılıklar", 2)
    for item in [
        "Neural heuristic optimalite garantisi taşımamaktadır.",
        "Ana deney sonuçları tek bir seed koşumuna dayanmaktadır.",
        "Tehdit paternleri önceden bilinen periyodik yapılardır.",
        "İrtifa, rüzgar, ayrıntılı yakıt modeli ve çoklu İHA kapsam dışındadır.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    column_break = doc.add_paragraph()
    column_break.paragraph_format.space_after = Pt(0)
    column_break.add_run().add_break(WD_BREAK.COLUMN)
    add_heading(doc, "6. Sonuç")
    add_body(doc, "Bu çalışmada MLP tabanlı öğrenilmiş bir heuristic, zaman-genişletmeli A* planlayıcısına entegre edilmiştir. Neural yaklaşım iki senaryoda Octile'dan daha az, iki senaryoda daha fazla düğüm genişletmiştir. Bulgular, öğrenilmiş heuristiclerin klasik yöntemlerin yerine her durumda geçemeyeceğini; ancak uygun çevre yapılarında arama verimliliği sağlayabileceğini göstermektedir. Gelecek çalışmalarda çoklu seed deneyleri, tam dinamik cost-to-go etiketleri ve admissible sınırlandırma yöntemleri incelenebilir.")

    add_heading(doc, "Kaynakça")
    refs = [
        "[1] P. E. Hart, N. J. Nilsson ve B. Raphael, “A formal basis for the heuristic determination of minimum cost paths,” IEEE TSSC, 1968.",
        "[2] S. Russell ve P. Norvig, Artificial Intelligence: A Modern Approach, 4. bs., Pearson, 2020.",
        "[3] D. P. Kingma ve J. Ba, “Adam: A method for stochastic optimization,” ICLR, 2015.",
        "[4] R. Yonetani vd., “Path Planning using Neural A* Search,” ICML, 2021.",
        "[5] S. M. LaValle, Planning Algorithms, Cambridge University Press, 2006.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(ref)
        set_run_font(run, size=7.5)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


def main():
    _, results, _, _ = report_data.run_experiments()
    path = build_article(results)
    print(f"Makale taslağı oluşturuldu: {path}")


if __name__ == "__main__":
    main()
