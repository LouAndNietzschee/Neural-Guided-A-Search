"""Create an editable English two-column article for the UAV project."""

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


SRC_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SRC_DIR.parent
REPORT_DIR = PROJECT_DIR / "report"
ASSET_DIR = REPORT_DIR / "assets_generated"
OUTPUT_DIR = Path("/workspace/output/docx")
OUTPUT_PATH = OUTPUT_DIR / "Neural_Guided_Astar_English_Article.docx"
sys.path.insert(0, str(SRC_DIR))

import generate_final_report as report_data  # noqa: E402
from generate_article_draft import (  # noqa: E402
    BLUE,
    MUTED,
    add_body,
    add_heading,
    add_picture,
    add_table,
    configure_footer,
    configure_section,
    configure_styles,
    set_run_font,
)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(
        "Neural-Guided A* Search for UAV Path Planning\n"
        "in Dynamic Threat Environments"
    )
    set_run_font(run, size=16, bold=True, color="17365D")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Kaan Arslan and Rukiye Narsu Oymak")
    set_run_font(run, size=10.5, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Biruni University, Artificial Intelligence Course Project\n"
        "Supervisor: Mahyar Teymournezhad"
    )
    set_run_font(run, size=8.5, italic=True, color=MUTED)


def add_abstract(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = p.add_run("Abstract—")
    set_run_font(lead, size=9, bold=True)
    text = (
        "This study presents a learned-heuristic A* framework for unmanned aerial vehicle "
        "(UAV) path planning in two-dimensional grids containing static obstacles, restricted "
        "areas, traversable risk zones, and periodic dynamic threats. A multilayer perceptron "
        "(MLP) estimates the remaining path cost from geometric and local environmental "
        "features, while a time-expanded A* planner represents each state as position and "
        "threat phase. The implementation is evaluated against Dijkstra search, Manhattan A*, "
        "and Octile A* in four operational scenarios. In the updated experiments, the learned "
        "heuristic expands 8.4% and 19.6% fewer nodes than Octile in the Simple Infiltration and "
        "Urban Operation scenarios, respectively, but expands 105.5% and 15.0% more nodes in "
        "the Corridor Passage and Dynamic Threat scenarios. Its route-cost deviations from the "
        "Dijkstra optimum are 0.90%, 2.03%, 5.35%, and 0.91%. These results show that learned "
        "guidance can improve search efficiency in selected map structures, but its benefit is "
        "scenario dependent and does not replace the optimality guarantees of an admissible "
        "heuristic. The method should therefore be viewed as a hybrid trade-off between search "
        "effort and route quality."
    )
    run = p.add_run(text)
    set_run_font(run, size=9)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = p.add_run("Keywords—")
    set_run_font(lead, size=8.5, bold=True)
    run = p.add_run(
        "A* search, Dijkstra, learned heuristic, multilayer perceptron, UAV path planning, "
        "dynamic threats, time-expanded planning."
    )
    set_run_font(run, size=8.5, italic=True)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_column_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.COLUMN)


def build_article(results):
    doc = Document()
    configure_styles(doc)

    first = doc.sections[0]
    configure_section(first, columns=1)
    configure_footer(first)
    doc.core_properties.title = (
        "Neural-Guided A* Search for UAV Path Planning in Dynamic Threat Environments"
    )
    doc.core_properties.author = "Kaan Arslan and Rukiye Narsu Oymak"
    doc.core_properties.subject = "Artificial Intelligence Course Project"

    add_title_block(doc)
    add_abstract(doc)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, columns=2)

    add_heading(doc, "1. Introduction")
    add_body(
        doc,
        "Safe and computationally efficient route planning is a central requirement for "
        "autonomous UAV missions. A feasible route cannot be selected only from the geometric "
        "distance between a start and a goal. Buildings, terrain, restricted zones, radar or "
        "surface-to-air missile regions, and changes in threat activity over time may alter both "
        "the validity and the cost of a route. The planning problem must therefore balance route "
        "quality, safety constraints, and computational effort."
    )
    add_body(
        doc,
        "A* is a standard graph-search method that combines the accumulated cost g(n) with a "
        "heuristic estimate h(n). Its efficiency depends strongly on the information contained in "
        "h(n). Geometric heuristics such as Manhattan and Octile distance are inexpensive, but "
        "they do not explicitly encode obstacle density or local risk. A learned heuristic can "
        "include this context and may guide the search toward promising regions earlier. However, "
        "a regression model is not automatically admissible and may overestimate the true "
        "remaining cost, weakening the optimality guarantee of classical A*."
    )
    add_body(
        doc,
        "This article combines the broader problem description of the original study with the "
        "corrected time-aware implementation and the latest experimental results. The principal "
        "contributions are:"
    )
    add_bullet(doc, "A time-expanded state representation for periodic threat activity.")
    add_bullet(doc, "An MLP heuristic using geometric, obstacle, direction, and risk features.")
    add_bullet(doc, "Directed reverse-Dijkstra labels consistent with asymmetric cell-entry costs.")
    add_bullet(doc, "A comparison of search effort and route quality across four scenarios.")

    add_heading(doc, "2. Related Work")
    add_body(
        doc,
        "Dijkstra's algorithm returns shortest paths for graphs with nonnegative edge costs, but "
        "it does not use goal-directed information and can expand a large part of the search "
        "space. A* improves this behavior by ranking states with f(n)=g(n)+h(n). If h(n) never "
        "exceeds the actual cost-to-go and the usual graph-search conditions hold, A* preserves "
        "optimality while reducing unnecessary exploration [1], [2]. Octile distance is a natural "
        "lower bound for an eight-connected grid with unit orthogonal moves and square-root-of-two "
        "diagonal moves. Manhattan distance is appropriate for four-connected movement, but it can "
        "overestimate in the eight-connected setting used here."
    )
    add_body(
        doc,
        "Learning-based planners attempt to exploit regularities that geometric formulas cannot "
        "represent. Previous work has learned cost maps, search priorities, or planning policies "
        "from demonstrations [5], [6]. The present project is narrower than a differentiable Neural "
        "A* system: it retains a conventional A* loop and replaces only the heuristic estimate with "
        "an MLP regressor. This distinction is important because the model affects node ordering, "
        "whereas collision checking, dynamic-state transitions, and path reconstruction remain "
        "explicit algorithmic components."
    )

    add_heading(doc, "3. Problem Formulation")
    add_heading(doc, "3.1 Grid Environment and Costs", 2)
    add_body(
        doc,
        "The environment is a 50 x 50 weighted grid. The UAV can move to eight neighboring cells. "
        "Orthogonal moves have a base cost of 1, and diagonal moves have a base cost of sqrt(2). "
        "Static obstacles, active threats, and prohibited cells are impassable. A passive threat "
        "cell remains traversable but receives a risk multiplier of 1.5. Consequently, edge costs "
        "are directional when the cost is charged upon entering the destination cell."
    )
    cell_rows = [
        ["Cell type", "Traversal", "Cost treatment"],
        ["Free space", "Allowed", "Base move cost"],
        ["Static obstacle", "Blocked", "Not applicable"],
        ["Active threat", "Blocked", "Not applicable"],
        ["Passive threat", "Allowed", "1.5 risk multiplier"],
        ["Restricted zone", "Blocked", "Not applicable"],
    ]
    add_table(doc, cell_rows, [1450, 1150, 2174], "Table 1. Cell semantics in the planning grid.")

    add_heading(doc, "3.2 Time-Expanded Dynamic State", 2)
    add_body(
        doc,
        "The dynamic planner represents a state as (r,c,t), where r and c are the grid coordinates "
        "and t is the discrete arrival time. A movement or wait action advances time by one step, "
        "and the target cell is checked against the threat configuration at the arrival time. To "
        "keep the state space finite, t is stored modulo the least common multiple of all threat "
        "periods. Thus, the same location reached in different threat phases remains a distinct "
        "search state. This prevents an apparently short spatial route from being accepted when a "
        "periodic threat is active at the corresponding arrival time."
    )

    add_heading(doc, "4. Learned Heuristic")
    add_heading(doc, "4.1 Feature Representation", 2)
    add_body(
        doc,
        "For each state-goal pair, the model receives eight features: Manhattan, Euclidean, and "
        "Octile distances; signed row and column offsets; local obstacle density; the obstacle "
        "ratio in the direction of the goal; and the risk value of the current cell. These features "
        "combine global geometric progress with a compact description of local traversability. The "
        "arrival-time phase is used when the heuristic is queried in dynamic search, so feature "
        "extraction is consistent with the current threat snapshot."
    )

    add_heading(doc, "4.2 Network and Training", 2)
    architecture_rows = [
        ["Stage", "Size", "Activation"],
        ["Input", "8 features", "-"],
        ["Hidden 1", "64 neurons", "ReLU"],
        ["Hidden 2", "64 neurons", "ReLU"],
        ["Hidden 3", "32 neurons", "ReLU"],
        ["Output", "1 value", "Linear"],
    ]
    add_table(doc, architecture_rows, [1450, 1500, 1824], "Table 2. MLP architecture.")
    add_body(
        doc,
        "The target is the remaining path cost obtained by Dijkstra search. Because cell-entry "
        "costs are directional, labels are generated by reverse Dijkstra using reversed edge "
        "semantics rather than by a naive symmetric-distance calculation. The final dataset "
        "contains 157,693 training and 31,668 validation samples. Training and validation maps use "
        "separate seed groups to reduce direct map leakage. The network is optimized for 80 epochs "
        "with Adam and mean squared error (MSE). Final training and validation MSE values are "
        "15.5985 and 11.9248, respectively."
    )
    add_picture(doc, ASSET_DIR / "training_curve.png", 2.55, "Figure 1. MLP training and validation loss.")

    add_heading(doc, "5. Experimental Setup")
    add_body(
        doc,
        "Four maps are used: Simple Infiltration, Urban Operation, Corridor Passage, and Dynamic "
        "Threat. The compared planners are Dijkstra, Manhattan A*, Octile A*, and Neural A*. "
        "Dijkstra provides the reference optimum. Octile is the principal admissible geometric "
        "baseline for the movement model. Manhattan is retained as an intentionally aggressive "
        "comparison, but its low expansion count must not be interpreted as an optimality guarantee "
        "in an eight-connected grid."
    )
    add_body(
        doc,
        "Two metrics are reported. Expanded nodes approximate the amount of graph exploration. "
        "Route cost measures solution quality and includes movement and risk penalties. Neural cost "
        "deviation is computed as 100(C_N-C_D)/C_D, where C_N and C_D are the Neural and Dijkstra "
        "route costs. The benchmark summarizes one fixed, reproducible run per scenario; it is not "
        "a randomized controlled trial or a multi-seed statistical study."
    )

    node_rows = [["Scenario", "Dij.", "Man.", "Oct.", "Neural"]]
    cost_rows = [["Scenario", "Dij.", "Man.", "Oct.", "Neural"]]
    name_map = {
        "Basit Sizma": "Simple",
        "Sehir Operasyonu": "Urban",
        "Koridor Gecisi": "Corridor",
        "Dinamik Tehdit": "Dynamic",
    }
    for scenario, title, _ in report_data.SCENARIOS:
        rs = results[scenario]
        node_rows.append([
            name_map[title],
            f"{rs[report_data.METHODS[0]].nodes_expanded:,}",
            f"{rs[report_data.METHODS[1]].nodes_expanded:,}",
            f"{rs[report_data.METHODS[2]].nodes_expanded:,}",
            f"{rs[report_data.METHODS[3]].nodes_expanded:,}",
        ])
        cost_rows.append([
            name_map[title],
            f"{rs[report_data.METHODS[0]].cost:.2f}",
            f"{rs[report_data.METHODS[1]].cost:.2f}",
            f"{rs[report_data.METHODS[2]].cost:.2f}",
            f"{rs[report_data.METHODS[3]].cost:.2f}",
        ])
    add_table(doc, node_rows, [1400, 844, 844, 844, 842], "Table 3. Expanded node counts.")
    add_table(doc, cost_rows, [1400, 844, 844, 844, 842], "Table 4. Route costs.")

    add_heading(doc, "6. Results")
    add_heading(doc, "6.1 Quantitative Comparison", 2)
    add_body(
        doc,
        "Relative to Octile, Neural A* expands 8.4% fewer nodes in Simple Infiltration and 19.6% "
        "fewer in Urban Operation. The direction reverses in the more constrained cases: Neural "
        "expands 105.5% more nodes in Corridor Passage and 15.0% more in Dynamic Threat. Neural "
        "route-cost deviations from the Dijkstra optimum are 0.90%, 2.03%, 5.35%, and 0.91% in the "
        "four scenarios. Octile matches the Dijkstra route cost in every reported scenario, whereas "
        "the learned heuristic accepts a small to moderate loss of route quality."
    )
    comparison_rows = [
        ["Scenario", "Neural vs. Octile nodes", "Neural cost deviation"],
        ["Simple", "-8.4%", "+0.90%"],
        ["Urban", "-19.6%", "+2.03%"],
        ["Corridor", "+105.5%", "+5.35%"],
        ["Dynamic", "+15.0%", "+0.91%"],
    ]
    add_table(doc, comparison_rows, [1400, 1800, 1574], "Table 5. Relative Neural performance.")

    add_heading(doc, "6.2 Route-Level Observations", 2)
    add_body(
        doc,
        "The route plots explain why a single heuristic is not uniformly dominant. In the simple "
        "and urban maps, environmental features help the model prefer openings around large blocked "
        "regions. In the narrow corridor, however, the topology already imposes a small number of "
        "viable directions. Octile then provides a stable geometric ordering, while the MLP's local "
        "estimation errors create additional exploration. In the dynamic map, a spatially attractive "
        "choice may become unavailable at a later phase, making accurate temporal cost-to-go "
        "prediction harder."
    )

    figures = [
        ("01_basit", "Figure 2. Routes in the Simple Infiltration scenario."),
        ("02_sehir", "Figure 3. Routes in the Urban Operation scenario."),
        ("03_koridor", "Figure 4. Routes in the Corridor Passage scenario."),
        ("04_dinamik", "Figure 5. Routes in the Dynamic Threat scenario."),
    ]
    for prefix, caption in figures:
        add_picture(doc, ASSET_DIR / f"{prefix}_routes.png", 2.58, caption)

    add_heading(doc, "6.3 Heuristic Surfaces", 2)
    add_body(
        doc,
        "The Manhattan surface changes smoothly with geometric distance and is insensitive to the "
        "map structure. The neural surface is less regular because local obstacles and risk features "
        "modify the predicted cost. This added structure can guide the search around difficult "
        "regions, but it can also introduce local overestimation or inconsistent ordering."
    )
    add_picture(doc, ASSET_DIR / "urban_manhattan_heatmap.png", 2.05, "Figure 6. Manhattan heuristic surface.")
    add_picture(doc, ASSET_DIR / "urban_neural_heatmap.png", 2.05, "Figure 7. Neural heuristic surface.")

    doc.add_page_break()
    add_heading(doc, "7. Discussion")
    add_body(
        doc,
        "The updated results support a more cautious conclusion than the original paper. Learned "
        "guidance is useful when the feature representation captures map regularities that are not "
        "visible to a geometric lower bound. This occurs in the simple and urban scenarios. It is "
        "not sufficient, however, to claim that Neural A* is globally faster or better. The corridor "
        "and dynamic results demonstrate that a compact MLP may misorder states when passage "
        "structure is highly constrained or when future threat phases dominate the remaining cost."
    )
    add_body(
        doc,
        "The comparison also separates node count from solution quality. Manhattan expands very few "
        "nodes, but this behavior partly follows from overaggressive estimation in an eight-connected "
        "grid. Similarly, a learned heuristic may reduce expansions by prioritizing one corridor "
        "strongly, yet return a more expensive path. Octile is therefore the most reliable overall "
        "baseline in this experiment: it preserves the reference route cost and remains competitive "
        "in search effort. Neural guidance is better interpreted as an optional acceleration strategy "
        "for map families on which its benefit has been validated."
    )
    add_body(
        doc,
        "A counterfactual route comparison is useful when presenting the system: selecting a direct "
        "geometric route can produce a short-looking path that intersects a blocked cell or reaches a "
        "periodic threat during its active phase. The time-expanded planner instead permits a detour "
        "or a wait action. The resulting route may be spatially longer but feasible at the actual "
        "arrival times. This distinction is central to interpreting dynamic path-planning results."
    )

    add_heading(doc, "8. Limitations and Future Work")
    add_bullet(doc, "The neural heuristic is not admissible and does not guarantee an optimal route.")
    add_bullet(doc, "The main benchmark uses one fixed run per scenario; multi-seed confidence intervals are absent.")
    add_bullet(doc, "Training targets are directional snapshot costs rather than full dynamic cost-to-go values.")
    add_bullet(doc, "The environment is a 2D grid and omits altitude, wind, detailed energy use, and vehicle dynamics.")
    add_bullet(doc, "Threat schedules are periodic and known in advance; sensing uncertainty is not modeled.")
    add_bullet(doc, "Per-node MLP inference overhead may offset node savings on small maps.")
    add_body(
        doc,
        "Future work should repeat the benchmark over multiple independently generated maps, report "
        "runtime as well as expansions, and train on labels computed in the full time-expanded state "
        "space. A safe hybrid could combine the model with Octile through bounded prediction, "
        "fallback rules, or an admissibility correction. Larger maps, three-dimensional motion, "
        "uncertain threat forecasts, and multi-UAV coordination would provide more realistic tests of "
        "generalization. Graph-based or convolutional encoders may also represent nonlocal obstacle "
        "structure more effectively than the current hand-crafted feature vector."
    )

    add_heading(doc, "9. Conclusion")
    add_body(
        doc,
        "This work integrates an MLP cost-to-go estimator with a time-expanded A* planner for UAV "
        "routing under static and periodic dynamic constraints. The revised implementation uses "
        "arrival-time-aware heuristic queries, directed reverse-Dijkstra labels, and separate "
        "training and validation seed groups. Across four scenarios, Neural A* outperforms Octile in "
        "expanded nodes twice and underperforms it twice, while producing route costs 0.90% to 5.35% "
        "above the Dijkstra optimum. The evidence therefore supports a scenario-dependent hybrid "
        "method, not a universal replacement for admissible geometric heuristics."
    )

    add_heading(doc, "References")
    refs = [
        "[1] P. E. Hart, N. J. Nilsson, and B. Raphael, 'A formal basis for the heuristic determination of minimum cost paths,' IEEE Transactions on Systems Science and Cybernetics, vol. 4, no. 2, pp. 100-107, 1968, doi: 10.1109/TSSC.1968.300136.",
        "[2] S. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, 4th ed. Pearson, 2020, ISBN: 978-0-13-461099-3.",
        "[3] S. M. LaValle, Planning Algorithms. Cambridge: Cambridge University Press, 2006, doi: 10.1017/CBO9780511546877.",
        "[4] D. P. Kingma and J. Ba, 'Adam: A method for stochastic optimization,' in Proc. 3rd International Conference on Learning Representations (ICLR), 2015, arXiv:1412.6980.",
        "[5] R. Yonetani, T. Taniai, M. Barekatain, M. Nishimura, and A. Kanezaki, 'Path Planning using Neural A* Search,' in Proceedings of the 38th International Conference on Machine Learning, PMLR, vol. 139, pp. 12029-12039, 2021. [Online]. Available: https://proceedings.mlr.press/v139/yonetani21a.html.",
        "[6] S. Choudhury, M. Bhardwaj, S. Arora, A. Kapoor, G. Ranade, S. Scherer, and D. Dey, 'Data-driven planning via imitation learning,' The International Journal of Robotics Research, vol. 37, no. 13-14, pp. 1632-1672, 2018, doi: 10.1177/0278364918781001.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(1.5)
        run = p.add_run(ref)
        set_run_font(run, size=7.5)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


def main():
    _, results, _, _ = report_data.run_experiments()
    path = build_article(results)
    print(f"English article created: {path}")


if __name__ == "__main__":
    main()
