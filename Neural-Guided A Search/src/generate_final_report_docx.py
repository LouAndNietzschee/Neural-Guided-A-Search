"""Create an editable Word version of the updated Neural-Guided A* report."""

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SRC_DIR.parent
REPORT_DIR = PROJECT_DIR / "report"
ASSET_DIR = REPORT_DIR / "assets_generated"
OUTPUT_PATH = REPORT_DIR / "Neural_Guided_Astar_Guncel_Rapor.docx"
sys.path.insert(0, str(SRC_DIR))

import generate_final_report as report_data  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "F4F6F9"
GRID = "9AA7B4"
MUTED = "5D6670"
TOTAL_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != TOTAL_DXA:
        raise ValueError(f"Table widths must sum to {TOTAL_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOTAL_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    paragraph.add_run("Sayfa ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet_style = doc.styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(11)
    bullet_style.paragraph_format.left_indent = Inches(0.375)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.194)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style.paragraph_format.line_spacing = 1.208

    if "Caption Custom" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("Caption Custom", 1)
    else:
        caption = doc.styles["Caption Custom"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)

    if "Code Custom" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("Code Custom", 1)
    else:
        code = doc.styles["Code Custom"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.left_indent = Inches(0.15)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Neural-Guided A* Search | Yapay Zeka Dersi Dönem Projesi")
    set_run_font(run, size=8, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    add_page_field(fp)
    for run in fp.runs:
        set_run_font(run, size=8, color=MUTED)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        lead.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_caption(doc, text):
    return doc.add_paragraph(text, style="Caption Custom")


def add_picture(doc, path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Custom")
    p.add_run(text)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F3F5")
    p_pr.append(shd)
    return p


def add_data_table(doc, rows, widths_dxa, caption=None, center_cols=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(value)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if center_cols and c_idx in center_cols:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=9, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
    set_repeat_header(table.rows[0])
    set_table_geometry(table, widths_dxa)
    if caption:
        add_caption(doc, caption)
    return table


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    logo = REPORT_DIR / "biruni_logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.35))
    for text, size, bold, color, after in [
        ("Neural-Guided A* Search", 25, True, DARK_BLUE, 8),
        ("Dinamik Tehdit Ortamlarında İHA Rota Planlaması için\nSinir Ağı Destekli A* Arama", 15, False, DARK_BLUE, 28),
        ("Yapay Zeka Dersi Dönem Projesi\nGüncellenmiş Proje Raporu", 13, True, "000000", 34),
        ("HAZIRLAYANLAR", 11, True, "000000", 4),
        ("Kaan Arslan (230404045)\nRukiye Narsu Oymak (230404001)", 11, False, "000000", 22),
        ("DANIŞMAN ÖĞRETMEN", 11, True, "000000", 4),
        ("Mahyar Teymournezhad", 11, False, "000000", 28),
        ("Haziran 2026", 11, True, DARK_BLUE, 0),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        for index, line in enumerate(text.split("\n")):
            if index:
                p.add_run().add_break()
            run = p.add_run(line)
            set_run_font(run, size=size, bold=bold, color=color)
    doc.add_page_break()


def build_document(results):
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "Neural-Guided A* Search - Güncel Proje Raporu"
    doc.core_properties.author = "Kaan Arslan ve Rukiye Narsu Oymak"
    doc.core_properties.subject = "Yapay Zeka Dersi Dönem Projesi"
    add_cover(doc)

    # Page 2
    add_heading(doc, "Proje Özeti")
    add_body(doc, "Bu çalışmada, statik engeller ve zamana bağlı tehditler içeren 50x50 grid ortamlarda İHA rota planlaması için MLP tabanlı öğrenilmiş bir heuristic ile zaman-genişletmeli A* birleştirilmiştir. Dijkstra, Manhattan, Octile ve Neural heuristic dört senaryoda düğüm genişletme ve rota maliyeti açısından karşılaştırılmıştır. Güncel uygulamada Neural heuristic her düğümün gerçek varış zamanını kullanır; eğitim etiketleri yönlü hücre maliyetlerine uygun ters Dijkstra ile üretilir ve doğrulama haritaları eğitimden farklı seed değerleriyle oluşturulur. Sonuçlar Neural yaklaşımın Basit Sızma ve Şehir Operasyonu senaryolarında Octile'dan daha az düğüm genişlettiğini; Koridor ve Dinamik Tehdit senaryolarında ise Octile'ın daha verimli kaldığını göstermiştir.")
    add_body(doc, "Anahtar Kelimeler: A* araması, Dijkstra, MLP, öğrenilmiş heuristic, İHA rota planlama, dinamik tehdit, zaman-genişletmeli arama.")
    add_heading(doc, "İçindekiler")
    add_bullets(doc, [
        "1. Giriş ve Katkılar", "2. Sistem Tasarımı", "3. Sinir Ağı Tabanlı Heuristic",
        "4. Uygulama Düzeltmeleri", "5. Deneysel Kurulum", "6. Deneysel Sonuçlar",
        "7. Senaryo Analizleri", "8. Heatmap Analizi", "9. Tartışma ve Sınırlılıklar",
        "10. Çalıştırma Talimatları ve Sonuç", "Kaynakça", "Ek A. Proje Yapısı ve Testler",
    ])
    doc.add_page_break()

    # Page 3
    add_heading(doc, "1. Giriş")
    add_body(doc, "Otonom İHA sistemlerinde güvenli rota planlama, yalnızca geometrik olarak kısa bir yol bulma problemi değildir. Radar, SAM ve elektronik harp bölgelerinin zamanla aktif veya pasif hale gelmesi, planlayıcının her hücreyi İHA'nın varış zamanında değerlendirmesini gerektirir. A* algoritması uygun bir heuristic ile arama alanını daraltabilir; ancak öğrenilmiş heuristic gerçek maliyeti aşarsa optimal çözüm garantisi kaybolur. Bu proje, bu performans-kalite takasını deneysel olarak incelemektedir.")
    add_heading(doc, "1.1 Çalışmanın Katkıları", 2)
    add_bullets(doc, [
        "Zaman-genişletmeli durum uzayı ile periyodik tehditlerin rota boyunca kontrol edilmesi.",
        "Dijkstra, Manhattan, Octile ve MLP heuristiclerinin aynı deney düzeninde karşılaştırılması.",
        "Neural özelliklerinin her düğümün gerçek varış zamanıyla hesaplanması.",
        "Yönlü hareket maliyetleri için doğru ters-Dijkstra ground-truth üretimi.",
        "Eğitim ve doğrulama haritalarının farklı seed gruplarından üretilmesi.",
        "Koyu temalı rota karşılaştırması, benchmark panelleri ve zaman uyumlu snapshot çıktıları.",
    ])
    add_heading(doc, "1.2 Kapsam", 2)
    add_body(doc, "Çalışma iki boyutlu, ayrık zamanlı ve periyodik tehdit bilgisi önceden bilinen bir simülasyonla sınırlıdır. Uçuş dinamiği, irtifa, rüzgar ve gerçek zamanlı algılama proje kapsamı dışındadır.")
    doc.add_page_break()

    # Page 4
    add_heading(doc, "2. Sistem Tasarımı")
    add_heading(doc, "2.1 Operasyon Ortamı", 2)
    add_body(doc, "Ortam 50x50 hücreli bir grid olarak modellenmiştir. İHA sekiz yönde hareket edebilir. Düz hareket maliyeti 1,0, çapraz hareket maliyeti sqrt(2)'dir. Pasif tehdit bölgesine giriş 1,5 risk çarpanıyla maliyetlendirilir.")
    add_data_table(doc, [
        ["Hücre Tipi", "Açıklama", "Geçilebilirlik / Maliyet"],
        ["Boş saha", "Açık hava sahası", "Evet, 1,0"],
        ["Statik engel", "Dağ veya bina", "Hayır"],
        ["Aktif tehdit", "Aktif radar / SAM", "Hayır"],
        ["Pasif tehdit", "Geçici olarak kapalı tehdit", "Evet, 1,5 risk çarpanı"],
        ["Yasak bölge", "Kritik tesis veya hava sahası", "Hayır"],
    ], [1800, 4080, 3480], "Tablo 1. Grid ortamında kullanılan hücre tipleri.")
    add_heading(doc, "2.2 Dinamik Durum Uzayı", 2)
    add_body(doc, "Dinamik A* durumu (r, c, t) olarak tutar. Her hareket veya bekleme eylemi zamanı bir adım ilerletir. Aynı konumun farklı tehdit fazlarındaki halleri farklı durumlar olarak ele alınır. Tehdit periyotlarının EKOK'u kullanılarak durum uzayı sonlu tutulur.")
    add_heading(doc, "2.3 Karşılaştırılan Yöntemler", 2)
    add_data_table(doc, [
        ["Yöntem", "Heuristic", "Optimalite"],
        ["Dijkstra", "h(n)=0", "Optimal referans"],
        ["Manhattan", "|dr|+|dc|", "8-yönlü harekette garanti yok"],
        ["Octile", "max+(sqrt(2)-1)min", "Admissible ve consistent"],
        ["Neural", "MLP(x(n,g,t))", "Admissibility garantisi yok"],
    ], [1800, 4080, 3480])
    doc.add_page_break()

    # Page 5
    add_heading(doc, "3. Sinir Ağı Tabanlı Heuristic")
    add_heading(doc, "3.1 Model Mimarisi", 2)
    add_body(doc, "NumPy ile sıfırdan uygulanan MLP, sekiz özellikten hedefe kalan maliyeti tahmin eden bir regresyon modelidir. ReLU gizli katmanları, lineer çıkış, MSE kaybı ve Adam optimizer kullanılır.")
    add_data_table(doc, [
        ["Katman", "Boyut", "Aktivasyon"], ["Giriş", "8", "-"],
        ["Gizli 1", "64", "ReLU"], ["Gizli 2", "64", "ReLU"],
        ["Gizli 3", "32", "ReLU"], ["Çıkış", "1", "Lineer"],
    ], [3000, 3000, 3360], "Tablo 2. MLP mimarisi.", center_cols={1, 2})
    add_heading(doc, "3.2 Giriş Özellikleri", 2)
    add_data_table(doc, [
        ["No", "Özellik", "Amaç"],
        ["1", "Manhattan mesafesi", "Normalize geometrik uzaklık"],
        ["2", "Euclidean mesafesi", "Kuş uçuşu uzaklık"],
        ["3", "Octile mesafesi", "8-yönlü teorik alt sınır"],
        ["4-5", "Satır ve sütun farkı", "Hedef yönü"],
        ["6", "5x5 lokal engel yoğunluğu", "Yerel çevre bilgisi"],
        ["7", "Hedef doğrultusu engel oranı", "Rota üzerindeki yapısal zorluk"],
        ["8", "Hücre riski", "Aktif/pasif tehdit bilgisi"],
    ], [900, 3900, 4560], "Tablo 3. Neural heuristic giriş özellikleri.", center_cols={0})
    doc.add_page_break()

    # Page 6
    add_heading(doc, "3.3 Eğitim ve Doğrulama")
    add_body(doc, "Son model 157.693 eğitim ve 31.668 doğrulama örneğiyle 80 epoch eğitilmiştir. Eğitim haritaları 42 tabanlı seed grubundan, doğrulama haritaları 10042 tabanlı ayrı seed grubundan üretilmiştir. Son eğitim MSE değeri 15,5985; doğrulama MSE değeri 11,9248'dir. Farklı haritaların kullanılması veri sızıntısı riskini azaltır.")
    add_picture(doc, ASSET_DIR / "training_curve.png", 6.25, "Şekil 1. MLP eğitim ve doğrulama kaybının epoch boyunca değişimi.")
    add_heading(doc, "4. Uygulama Düzeltmeleri")
    add_data_table(doc, [
        ["Düzeltme", "Güncel Uygulama"],
        ["Zaman farkındalığı", "Neural özellikler her düğümün varış zamanı ile hesaplanır."],
        ["Ground-truth yönü", "Girilen hücreye bağlı maliyet için ters kenar maliyeti kullanılır."],
        ["Doğrulama ayrımı", "Eğitim ve validation farklı harita seed'lerinden gelir."],
        ["Arama sınırı", "max_nodes limiti gerçek genişletme sayısına uygulanır."],
        ["Sunum", "Koyu tema, benchmark ve time_path tabanlı snapshot eklenmiştir."],
    ], [2700, 6660], "Tablo 4. Birleşik sürümde yapılan temel düzeltmeler.")
    doc.add_page_break()

    # Page 7
    add_heading(doc, "5. Deneysel Kurulum")
    add_body(doc, "Dört senaryo aynı seed=42 haritalarında ve başlangıç-hedef çiftlerinde çalıştırılmıştır. Dijkstra optimal referans, Octile güçlü admissible baseline, Manhattan ise 8-yönlü modelde inadmissible karşılaştırma olarak kullanılmıştır. Ana metrikler genişletilen düğüm ve rota maliyetidir.")
    add_heading(doc, "6. Toplu Deney Sonuçları")
    node_rows = [["Senaryo", "Dijkstra", "Manhattan", "Octile", "Neural", "N vs O"]]
    cost_rows = [["Senaryo", "Dijkstra", "Manhattan", "Octile", "Neural", "Sapma"]]
    for scenario, title, _ in report_data.SCENARIOS:
        rs = results[scenario]
        node_rows.append([
            title, f"{rs[report_data.METHODS[0]].nodes_expanded:,}".replace(",", "."),
            f"{rs[report_data.METHODS[1]].nodes_expanded:,}".replace(",", "."),
            f"{rs[report_data.METHODS[2]].nodes_expanded:,}".replace(",", "."),
            f"{rs[report_data.METHODS[3]].nodes_expanded:,}".replace(",", "."),
            f"{report_data.pct(rs[report_data.METHODS[2]].nodes_expanded, rs[report_data.METHODS[3]].nodes_expanded):+.1f}%",
        ])
        cost_rows.append([
            title, f"{rs[report_data.METHODS[0]].cost:.2f}", f"{rs[report_data.METHODS[1]].cost:.2f}",
            f"{rs[report_data.METHODS[2]].cost:.2f}", f"{rs[report_data.METHODS[3]].cost:.2f}",
            f"{report_data.pct(rs[report_data.METHODS[0]].cost, rs[report_data.METHODS[3]].cost):+.2f}%",
        ])
    add_data_table(doc, node_rows, [2160, 1440, 1440, 1440, 1440, 1440], "Tablo 5. Genişletilen düğüm sayıları.", center_cols={1, 2, 3, 4, 5})
    add_data_table(doc, cost_rows, [2160, 1440, 1440, 1440, 1440, 1440], "Tablo 6. Rota maliyetleri ve Neural sapması.", center_cols={1, 2, 3, 4, 5})
    add_body(doc, "Neural, Basit Sızma'da Octile'a göre %8,4 ve Şehir Operasyonu'nda %19,6 daha az düğüm genişletmiştir. Buna karşılık Koridor'da %105,5, Dinamik Tehdit'te %15,0 daha fazla düğüm genişletmiştir. Bu sonuç kazancın senaryo yapısına bağlı olduğunu gösterir.")
    doc.add_page_break()

    # Pages 8-11
    notes = {
        "simple": "Neural, Octile'dan daha az düğüm genişletmiş; rota maliyeti Dijkstra optimumundan yaklaşık %0,91 sapmıştır.",
        "urban": "Neural, Octile'a göre %19,6 daha az düğüm genişletmiştir. Maliyet sapması yaklaşık %2,03'tür.",
        "corridor": "Dar geçit yapısında Octile en verimli güvenli baseline olmuştur. Neural daha fazla düğüm açmış ve maliyette yaklaşık %5,36 sapmıştır.",
        "dynamic": "Dinamik tehdit senaryosunda Octile daha az düğüm açmıştır. Neural maliyet sapması yaklaşık %0,92'dir.",
    }
    for index, (scenario, title, prefix) in enumerate(report_data.SCENARIOS, start=1):
        add_heading(doc, f"7.{index} {title} Senaryosu")
        add_body(doc, notes[scenario])
        add_picture(doc, ASSET_DIR / f"{prefix}_routes.png", 6.25, f"Şekil {index * 2}. {title} için zaman-genişletmeli rota karşılaştırması.")
        add_picture(doc, ASSET_DIR / f"{prefix}_benchmark.png", 6.25, f"Şekil {index * 2 + 1}. {title} için düğüm ve rota maliyeti paneli.")
        doc.add_page_break()

    # Page 12
    add_heading(doc, "8. Heuristic Heatmap Analizi")
    add_body(doc, "Heatmap'ler aynı Şehir Operasyonu haritasında Manhattan ve Neural heuristic değerlerinin uzamsal dağılımını gösterir. Manhattan yalnızca geometrik uzaklığa bağlıdır. Neural yüzey ise mesafe özelliklerine ek olarak lokal engel yoğunluğu, hedef doğrultusu ve t=0 anındaki hücre riskini kullanır. Renkler her panel içinde normalize edildiği için mutlak renk değil, uzamsal desen karşılaştırılmalıdır.")
    add_picture(doc, ASSET_DIR / "urban_manhattan_heatmap.png", 2.2, "Şekil 10. Şehir senaryosu Manhattan heuristic yüzeyi.")
    add_picture(doc, ASSET_DIR / "urban_neural_heatmap.png", 2.2, "Şekil 11. Şehir senaryosu Neural heuristic yüzeyi.")
    add_heading(doc, "8.1 Zaman Uyumlu Snapshot'lar", 2)
    add_body(doc, "Demo modu Neural rota için her beş adımda bir snapshot üretir. Her karede tehditlerin aktif/pasif durumu result.time_path alanındaki gerçek varış zamanı ile çizilir.")
    doc.add_page_break()

    # Page 13
    add_heading(doc, "9. Tartışma ve Sınırlılıklar")
    add_heading(doc, "9.1 Bulguların Yorumu", 2)
    add_body(doc, "Sonuçlar Neural heuristicin her ortamda klasik Octile'ı geçmediğini açıkça göstermektedir. Şehir ve basit haritalarda öğrenilen çevre özellikleri aramayı daha doğrudan yönlendirirken, dar koridor yapısında modelin tahmin hatası hem düğüm sayısını hem rota maliyetini artırmıştır. Neural yöntem, senaryoya bağlı bir performans-kalite takası olarak yorumlanmalıdır.")
    add_heading(doc, "9.2 Manhattan Sonuçlarının Yorumu", 2)
    add_body(doc, "Manhattan çok az düğüm açsa da 8-yönlü ve sqrt(2) çapraz maliyetli modelde admissible değildir. Bu nedenle düğüm sayısı tek başına başarı kanıtı sayılmaz.")
    add_heading(doc, "9.3 Sınırlılıklar", 2)
    add_bullets(doc, [
        "Neural heuristic admissible değildir ve optimalite garantisi vermez.",
        "Eğitim etiketleri belirli zaman fazındaki statik maliyet yüzeyinden üretilir; tam dinamik cost-to-go etiketi değildir.",
        "Ana deney tablosu tek seed koşumudur; istatistiksel sonuç için çoklu seed gerekir.",
        "MLP inference maliyeti, düğüm azalması olsa bile çalışma süresini artırabilir.",
        "Simülasyon 2D grid ve bilinen periyodik tehditlerle sınırlıdır.",
    ])
    add_heading(doc, "9.4 Gelecek Çalışmalar", 2)
    add_bullets(doc, [
        "Farklı seed'lerle çoklu tekrar ve güven aralığı raporlamak.",
        "Tam zaman-genişletmeli cost-to-go etiketleriyle eğitim yapmak.",
        "h=min(MLP, Octile) gibi admissible sınırlandırmaları karşılaştırmak.",
        "3D rota, yakıt, rüzgar ve çoklu İHA koordinasyonu eklemek.",
    ])
    doc.add_page_break()

    # Page 14
    add_heading(doc, "10. Çalıştırma Talimatları")
    add_body(doc, "Windows PowerShell:")
    add_code(doc, ".\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt")
    add_code(doc, ".\\.venv\\Scripts\\python.exe src\\main.py --mode demo --planner dynamic")
    add_code(doc, ".\\.venv\\Scripts\\python.exe src\\generate_final_report_docx.py")
    add_body(doc, "Test komutu:")
    add_code(doc, "python -m unittest discover -s tests -v")
    add_heading(doc, "11. Sonuç")
    add_body(doc, "Proje, klasik A* ile öğrenilmiş heuristic yaklaşımını dinamik İHA rota planlama problemi üzerinde bir araya getirmiştir. Düzeltilmiş uygulama zaman bilgisini Neural özelliklere taşımış, yönlü ground-truth maliyetini düzeltmiş ve doğrulamayı ayrı harita seed'leriyle yapmıştır. Neural heuristic iki senaryoda Octile'a göre daha az, iki senaryoda daha fazla düğüm genişletmiştir. Bu dengeli sonuç, öğrenilmiş heuristiclerin faydasının ortam yapısına ve tahmin kalitesine bağlı olduğunu göstermektedir.")
    add_heading(doc, "Kaynakça")
    refs = [
        "Hart, P. E., Nilsson, N. J., & Raphael, B. (1968). A formal basis for the heuristic determination of minimum cost paths. IEEE TSSC, 4(2), 100-107.",
        "Russell, S., & Norvig, P. (2020). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.",
        "Kingma, D. P., & Ba, J. (2015). Adam: A method for stochastic optimization. ICLR.",
        "Yonetani, R. et al. (2021). Path Planning using Neural A* Search. ICML.",
        "LaValle, S. M. (2006). Planning Algorithms. Cambridge University Press.",
    ]
    for index, ref in enumerate(refs, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)
    doc.add_page_break()

    # Page 15
    add_heading(doc, "Ek A. Proje Klasör Yapısı")
    add_code(doc, """neural_astar_uav_birlesik/
|-- src/
|   |-- environment.py
|   |-- astar.py
|   |-- neural_heuristic.py
|   |-- visualization.py
|   |-- main.py
|   |-- generate_final_report.py
|   `-- generate_final_report_docx.py
|-- data/mlp_heuristic.pkl
|-- tests/test_core.py
|-- visualizations/
|-- report/
|-- MODEL_CARD.md
|-- README.md
`-- requirements.txt""")
    add_heading(doc, "Ek B. Doğrulama Testleri")
    add_data_table(doc, [
        ["Test", "Beklenen Davranış", "Sonuç"],
        ["Ground-truth", "Ters etiket ileri Dijkstra maliyetiyle eşit", "Geçti"],
        ["Dinamik zaman", "Heuristic birden fazla varış zamanı alır", "Geçti"],
        ["Rota güvenliği", "Her konum kendi varış zamanında geçilebilir", "Geçti"],
        ["Corner-cutting", "İki engel arasından çapraz geçiş yok", "Geçti"],
        ["Neural cache", "Önbellek zaman adımını anahtara dahil eder", "Geçti"],
    ], [2300, 5660, 1400], "Tablo 7. Çekirdek doğrulama testleri.", center_cols={2})
    add_body(doc, "Beş testin tamamı başarıyla geçmiştir. Dinamik ve statik demo modları eğitilmiş modelle uçtan uca çalıştırılmıştır. Rapor tabloları güncel model ve seed=42 deney koşumundan üretilmiştir.")

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


def main():
    _, results, _, _ = report_data.run_experiments()
    path = build_document(results)
    print(f"Word raporu oluşturuldu: {path}")


if __name__ == "__main__":
    main()
